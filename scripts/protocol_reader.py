import os
from datetime import datetime

import pandas as pd
from django.core.files import File
from django.db import transaction
from docx.api import Document

from darcydb.darcy_app.models import (
    ChemCodes,
    DictDocOrganizations,
    DictEntities,
    Documents,
    DocumentsPath,
    Wells,
    WellsSample,
)

LABORATORY = 'ООО "ГИЦ ПВ"'
PARAMETERS_DICT = {
    "Алюминий, мг/дм3": 3403,
    "Железо общее, мг/дм3": 2409,
    "Марганец, мг/дм3": 2414,
    "Кадмий, мг/дм3": 3331,
    "Медь, мг/дм3": 3330,
    "Мышьяк, мг/дм3": 3340,
    "Ртуть, мг/дм3": 3337,
    "Свинец, мг/дм3": 3332,
    "Кальций, мг/дм3": 2405,
    "Магний, мг/дм3": 2404,
    "Натрий, мг/дм3": 2401,
    "Калий, мг/дм3": 2402,
    "Нитраты, мг/дм3": 2309,
    "Нитриты, мг/дм3": 2308,
    "Щелочность общая, ммоль-экв/дм3": 2209,
    "Гидрокарбонаты, мг/дм3": 2303,
    "Жесткость общая, °Ж": 2210,
    "Водородный показатель (рН), ед. pH": 2201,
    "Нефтепродукты (суммарно) , мг/дм3": 4003,
    "Мутность, ЕМФ": 2607,
    "Цветность, градусы": 2009,
    "Привкус, баллы": 2004,
    "Запах, баллы": 2002,
    "Перманганатная окисляемость, мгО/дм3": 2203,
    "Ион аммония, мг/дм3": 2406,
    "Сульфаты, мг/дм3": 2302,
    "Хлориды, мг/дм3": 2301,
    "Фториды, мг/дм3": 3303,
    "Сероводород, мг/дм3": 3504,
    "Общая минерализация (сухой остаток), мг/дм3": 2208,
    "Никель, мг/дм3": 3333,
    "Селен, мг/дм3": 3341,
    "Молибден, мг/дм3": 3338,
}


def run():
    customer = DictDocOrganizations.objects.filter(name="ООО «Дарси»").first()
    if not customer:
        customer = DictDocOrganizations(name="ООО «Дарси»")
        customer.save()
    executor = DictDocOrganizations.objects.filter(name=LABORATORY).first()
    if not executor:
        executor = DictDocOrganizations(name=LABORATORY)
        executor.save()

    wells = Wells.objects.only("pk", "extra")
    wells_dict = {well.extra["name_gwk"]: well for well in wells}
    chem_codes = ChemCodes.objects.only("chem_id")
    codes_dict = {chem_code.chem_id: chem_code for chem_code in chem_codes}

    root_path = "./scripts/protocols"
    files = os.listdir(root_path)
    wells = pd.read_csv(os.path.join(root_path, "wells.csv"))
    for file in files:
        if file.endswith(".docx"):
            doc = Document(os.path.join(root_path, file))
            keys = tuple(cell.text for cell in doc.tables[0].rows[0].cells)
            well = doc.paragraphs[25].text.split(",")[-1].strip("скважина №").replace("x", "")
            date = doc.paragraphs[24].text.split(":")[-1].strip("x ")
            num_prot = doc.paragraphs[18].text
            date_doc = doc.paragraphs[27].text.split("-")[-1].strip()
            if "Дата отбора" in well:
                well = doc.paragraphs[26].text.split(",")[-1].strip("скважина №")
                date = doc.paragraphs[25].text.split(":")[-1].strip("x ")
                date_doc = doc.paragraphs[28].text.split("-")[-1].strip()
                num_prot = doc.paragraphs[19].text
            date = datetime.strptime(date, "%d.%m.%Y")
            date = date.strftime("%Y-%m-%d")
            date_doc = datetime.strptime(date_doc, "%d.%m.%Y")
            date_doc = date_doc.strftime("%Y-%m-%d")
            name_gwk = wells[wells["№"] == well.replace("x", "")]["№ГВК"]

            if not name_gwk.empty:
                name_gwk = name_gwk.iloc[0]
                if wells_dict.get(name_gwk):
                    print(name_gwk)
                    with transaction.atomic():
                        doc_inst = Documents(
                            name=num_prot,
                            typo=DictEntities.objects.get(
                                entity__name="тип документа", name="Протокол анализа пробы воды"
                            ),
                            source=executor,
                            org_executor=executor,
                            org_customer=customer,
                            creation_date=date_doc,
                        )
                        doc_inst.save()
                        with open(os.path.join(root_path, file), "rb") as f:
                            word_file = File(f)
                            document_path = DocumentsPath(doc=doc_inst)
                            document_path.path.save(file, word_file)
                            document_path.save()
                        sample = WellsSample(well=wells_dict.get(name_gwk), date=date, name=num_prot, doc=doc_inst)
                        sample.save()
                        for row in doc.tables[0].rows[1:]:
                            parameter = (cell.text for cell in row.cells)
                            data = dict(zip(keys, parameter))
                            code = PARAMETERS_DICT.get(data["Номенклатура показателей, \nединицы измерения"])
                            values = float(data["Значение показателя"].replace("<", "").replace(">", "").split("±")[0])
                            sample.chemvalues.create(parameter=codes_dict.get(code), chem_value=values)
                else:
                    print(f"Not found: {name_gwk}")
